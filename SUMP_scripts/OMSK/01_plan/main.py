import os
import chardet
from docx import Document
from typing import Dict, List, Optional, Set, Tuple

FILE_EXAMPLE = "plan_example.txt"
##FILE_IN = "in_plan"
FILE_OUT = "plan_my"
INPUT_FILES = ["Приложение_5__перестановки_БВ_.docx", "Приложение_7__выгрузка_АЗ_.docx"]


class MyParser:
    def __init__(self) -> None:
        self.number_cycle: str = None
        self.name_TVS: str = None
        self.cell_from_bridge: str = ""
        self.cell_from_trolley: str = ""
        self.cell_to_bridge: str = ""
        self.cell_to_trolley: str = ""
        self.name_PS_SUZ: str = None
        self.cycle: int = 0
        self.cargo: int = 0   # тип цикла из list.txt (#DefCargo)
        
        self.status: str = "N"
        self.data_from: str = "0:00"
        self.time_from: str = "0:00"
        self.data_to: str = "0:00"
        self.time_to: str = "0:00"
        
        # Определение кодировки в исходном файле
        self.encoding = self._detect_encoding(FILE_EXAMPLE)
        print(f"encoding {FILE_EXAMPLE}: {self.encoding}")

    def _detect_encoding(self, file_path: str) -> str:
        """Определяет кодировку файла."""
##        ResultDict = Dict[str, Optional[str] | float]   # не поддерживается в знерщт 3.8.10
        with open(file_path, "rb") as f:
            raw_data = f.read(20000)
            result: ResultDict = chardet.detect(raw_data)  # Аннотация с использованием TypedDic
            encoding: Optional[str] = result.get("encoding")
            
            if encoding is None:
                raise ValueError("Не удалось определить кодировку файла.")
            return encoding

    def parse_docx_table(self, input_file: str) -> None:
        
        self.data_to_plan: Dict[int, Dict[str, str]] = {}  # Словарь для хранения данных
        
        # Проверка существования входного файла
        if not os.path.exists(input_file):
            print(f"Файл РГП {input_file} не найден!")
            return

        try:
            doc = Document(input_file)
            if not doc.tables:
                print(f"В документе {input_file} нет таблиц!")
                return
                
            for table in doc.tables:
                for row in table.rows:
                    _data = [cell.text.strip() for cell in row.cells]
##                    print(_data)

                    if "Выгрузка ТВС из АЗ в БВ" in _data:
                        self.cycle = 5
##                        print(f"Тип цикла {self.cycle=}")
                        break

                    if "Перестановки ОТВС в БВ" in _data:
                        self.cycle = 7
##                        print(f"Тип цикла {self.cycle=}")
                        break
                    
                    if len(_data) > 1 and _data[0].isdigit():
                        self._extract_data_from_row(_data)
                        print(_data)
                        
                    if self.number_cycle and self.number_cycle not in self.data_to_plan:
                        self._initialize_cycle_data()
        except Exception as e:
            print(f"Ошибка при парсинге файла {input_file}: {e}")

    def _extract_data_from_row(self, data: List) -> None:
        """Извлекает данные из строки таблицы."""
        (
            self.number_cycle,
            cell_from,
            _,
            self.name_TVS,
            self.name_PS_SUZ,
            _,
            _,
            _,
            cell_to,
            _,
            _,
            _,
            _,
        ) = data
        
        self.cell_from_bridge, self.cell_from_trolley = cell_from.split("-")
        self.cell_to_bridge, self.cell_to_trolley = cell_to.split("-")
        if (self.name_TVS and self.name_PS_SUZ):
            self.cargo = 301
        else:
            self.cargo = 300

    def _initialize_cycle_data(self):
        """Инициализирует данные для каждого цикла перегрузки."""
        self.data_to_plan[self.number_cycle] = {
            "name_TVS": self.name_TVS,
            "DefCycle": self.cycle,
            "DefCargo": self.cargo,
            "cell_from_bridge": int(self.cell_from_bridge),
            "cell_from_trolley": int(self.cell_from_trolley),
            "cell_to_bridge": int(self.cell_to_bridge),
            "cell_to_trolley": int(self.cell_to_trolley),
            "name_PS_SUZ": self.name_PS_SUZ,
            "status": self.status,
            "data_from": self.data_from,
            "time_from": self.time_from,
            "data_to": self.data_to,
            "time_to": self.time_to,       
        }
##        print(self.data_to_plan[self.number_cycle])

    def check_data(self):
        for tvs, details in self.data.items():
            print(details)
        return f"Не соответствие DefCargo и name_TVS {name_PS_SUZ=}"

    def check_tvs_name(self):
        tvs_names = set()
        for details in self.data.values():
            tvs_name = details.get("name_TVS")
            if tvs_name in tvs_names:
                return f"ТВС '{tvs_name}' повторяется."
            tvs_names.add(tvs_name)
        return f"Все {len(tvs_names)} ТВС имеют уникальные имена."

    def create_plan(self):
        # Парсинг DOCX-файлов
        for index, input_file in enumerate(INPUT_FILES):
            self.parse_docx_table(input_file)
##            print(FILE_OUT + str(index) + ".txt")
            self._create_plan(FILE_OUT + str(index) + ".txt")

    def _create_plan(self, out_file: str) -> None:
        # Форматы полей: (ключ, ширина пробелов, значение по умолчанию)
        field_formats = [
            ("index", 4, ""),
            ("DefCycle", 6, ""),
            ("DefCargo", 5, ""),
            ("name_TVS", 20, ""),
            ("cell_from_bridge", 6, ""),
            ("cell_from_trolley", 6, ""),
            ("cell_to_bridge", 6, ""),
            ("cell_to_trolley", 6, ""),
            ("status", 2, ""),
            ("data_from", 8, ""),
            ("time_from", 8, ""),
            ("data_to", 8, ""),
            ("time_to", 8, ""),
            ("name_PS_SUZ", 8, ""),
        ]

        try:
            with open(out_file, "w", encoding=self.encoding) as f_out:
                for idx, cycle in self.data_to_plan.items():
##                    print(idx, cycle)
                    fields = []
                    for key, width, default in field_formats:
                        value = (
                            str(idx) if key == "index" else cycle.get(key, default)
                        )
##                        value = cycle.get(key, default)
                        fields.append((value, width))
                    print(fields)
                    formatted_line = (
                        "".join(f"{value:>{width}}" for value, width in fields) + "\n"
                    )
                    print(formatted_line)
                    f_out.write(formatted_line)
                print(f"Файл {out_file} успешно создан.")
        except Exception as e:
            print(f"Ошибка при создании файла {out_file}: {e}")

    def get_data(self):
        return self.data_to_plan


##    def __str__(self):
##        self.result = ""
##        for name_subzone, v in self.data.items():
##            self.result += name_subzone
##            self.result += "\n"
##            self.result += "".join(f"{a:>20}:{b:>6}\n" for a, b in v.items())
##        return self.result


def main():
    my_parser: MyParser = MyParser()
##    print(my_parser.check_tvs_name())
    my_parser.create_plan()


if __name__ == "__main__":
    main()
